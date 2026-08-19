import csv
import io
import json
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


@dataclass(frozen=True)
class ParsedDocument:
    title: str
    content: str
    source_type: str
    metadata: dict[str, Any] = field(default_factory=dict)


class DocumentParser:
    """Parse common office formats into auditable plain text without executing macros."""

    supported_extensions = {".txt", ".md", ".json", ".csv", ".docx", ".xlsx", ".pdf"}

    def parse(self, path: str | Path) -> ParsedDocument:
        source = Path(path)
        suffix = source.suffix.lower()
        if suffix not in self.supported_extensions:
            raise ValueError(f"unsupported document format: {suffix or '<none>'}")
        content, details = getattr(self, f"_parse_{suffix[1:]}")(source)
        normalized = "\n".join(line.rstrip() for line in content.splitlines()).strip()
        if not normalized:
            raise ValueError("document contains no extractable text")
        return ParsedDocument(
            title=source.stem, content=normalized, source_type=suffix[1:],
            metadata={"file_name": source.name, "extension": suffix, **details},
        )

    @staticmethod
    def _parse_txt(path: Path):
        return path.read_text(encoding="utf-8-sig"), {}

    _parse_md = _parse_txt

    @staticmethod
    def _parse_json(path: Path):
        value = json.loads(path.read_text(encoding="utf-8-sig"))
        return json.dumps(value, ensure_ascii=False, indent=2), {"root_type": type(value).__name__}

    @staticmethod
    def _parse_csv(path: Path):
        rows = list(csv.reader(io.StringIO(path.read_text(encoding="utf-8-sig"))))
        return "\n".join(" | ".join(cell.strip() for cell in row) for row in rows), {"row_count": len(rows)}

    @staticmethod
    def _parse_docx(path: Path):
        from docx import Document

        document = Document(path)
        blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            blocks.extend(" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows)
        return "\n".join(blocks), {"paragraph_count": len(document.paragraphs), "table_count": len(document.tables)}

    @staticmethod
    def _parse_xlsx(path: Path):
        from openpyxl import load_workbook

        workbook = load_workbook(path, read_only=True, data_only=True)
        blocks: list[str] = []
        row_count = 0
        try:
            for sheet in workbook.worksheets:
                blocks.append(f"工作表：{sheet.title}")
                for row in sheet.iter_rows(values_only=True):
                    values = [str(value).strip() if value is not None else "" for value in row]
                    if any(values):
                        blocks.append(" | ".join(values))
                        row_count += 1
        finally:
            workbook.close()
        return "\n".join(blocks), {"sheet_count": len(workbook.sheetnames), "row_count": row_count}

    @staticmethod
    def _parse_pdf(path: Path):
        from pypdf import PdfReader

        reader = PdfReader(path)
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
        return "\n\n".join(page for page in pages if page), {"page_count": len(reader.pages)}
